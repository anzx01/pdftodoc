"""DOCX 生成：把每页 OCR 文本行写入 Word 文档（版式简化，逐页分页）。"""

import logging
from collections.abc import Sequence
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.section import Section
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell

from pdftodoc.core.ocr import BBox, OcrLine, OcrPage, OcrTable, PageImage, SealImage

logger = logging.getLogger(__name__)

_BODY_FONT = "SimSun"
_TABLE_CELL_MARGIN = 36


def build_docx(pages: Sequence[OcrPage], dst: Path) -> None:
    """按页顺序写入文本行，页间插入分页符。"""
    dst.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for idx, page in enumerate(pages):
        _configure_ocr_page(document.sections[-1], page)
        if page.text_lines or page.tables or page.seals:
            _write_layout_page(document, page)
        else:
            for line in page.lines:
                document.add_paragraph(line)
        if not page.lines and not page.tables:
            document.add_paragraph("")  # 空页占位，保持页数对应
        if idx < len(pages) - 1:
            document.add_page_break()  # type: ignore[no-untyped-call]  # python-docx 缺类型标注
    document.save(str(dst))
    logger.info("已生成 DOCX：%s（%d 页）", dst, len(pages))


def build_image_docx(pages: Sequence[PageImage], dst: Path) -> None:
    """把扫描页原图逐页写入 Word，优先保留版式、公章与原文。"""
    dst.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for idx, page in enumerate(pages):
        section = document.sections[-1] if idx == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        _configure_page(section, page)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(BytesIO(page.png), width=Pt(page.width_pt))
    document.save(str(dst))
    logger.info("已生成版式优先 DOCX：%s（%d 页）", dst, len(pages))


def _configure_page(section: Section, page: PageImage) -> None:
    section.page_width = Pt(page.width_pt)
    section.page_height = Pt(page.height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _append_vml_seal(run: object, page: OcrPage, seal: SealImage, index: int) -> None:
    if not seal.png or not page.image_width_px or not page.image_height_px:
        return
    x1, y1, x2, y2 = seal.bbox
    left = x1 * _scale_x(page)
    top = y1 * _scale_y(page)
    width = max(1.0, (x2 - x1) * _scale_x(page))
    height = max(1.0, (y2 - y1) * _scale_y(page))
    r_id, _ = run.part.get_or_add_image(BytesIO(seal.png))  # type: ignore[attr-defined]
    xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <v:shape id="seal_{page.page_index}_{index}" type="#_x0000_t75"
        style="position:absolute;margin-left:{left:.2f}pt;margin-top:{top:.2f}pt;width:{width:.2f}pt;height:{height:.2f}pt;z-index:251660288;mso-position-horizontal-relative:page;mso-position-vertical-relative:page;mso-wrap-style:none"
        stroked="f" filled="f" o:allowincell="f">
        <v:imagedata r:id="{r_id}" o:title="seal"/>
      </v:shape>
    </w:pict>
    """
    run._r.append(parse_xml(xml))  # type: ignore[attr-defined]


def _write_layout_page(document: DocxDocument, page: OcrPage) -> None:
    tables = sorted(page.tables, key=lambda item: (item.bbox[1], item.bbox[0]))
    outside_lines = [
        line
        for line in sorted(page.text_lines, key=_line_sort_key)
        if not any(_contains_line(table.bbox, line) for table in tables)
    ]

    cursor = 0
    cursor_y = 0
    for table in tables:
        before, cursor = _lines_before(outside_lines, cursor, table.bbox[1])
        for line in before:
            cursor_y = _add_ocr_line(document, line, page, cursor_y)
        cursor_y = _add_table(document, table, page, cursor_y)

    for line in outside_lines[cursor:]:
        cursor_y = _add_ocr_line(document, line, page, cursor_y)

    for index, seal in enumerate(
        sorted(page.seals, key=lambda item: (item.bbox[1], item.bbox[0]))
    ):
        _add_seal(document, seal.png, seal.bbox, page, index)


def _lines_before(lines: Sequence[OcrLine], start: int, top: int) -> tuple[list[OcrLine], int]:
    before: list[OcrLine] = []
    index = start
    while index < len(lines):
        line = lines[index]
        if line.box is not None and _center(line.box)[1] >= top:
            break
        before.append(line)
        index += 1
    return before, index


def _add_table(
    document: DocxDocument, source: OcrTable, page: OcrPage, cursor_y: int
) -> int:
    _add_vertical_gap(document, page, cursor_y, source.bbox[1])
    table = document.add_table(rows=len(source.rows), cols=source.column_count)
    table.style = "Table Grid"
    row_heights = _table_row_heights(source)
    _set_table_geometry(table, source, page, row_heights)
    for row_index, row in enumerate(source.rows):
        col_index = 0
        for source_cell in row:
            span = max(1, source_cell.col_span)
            end_index = min(source.column_count - 1, col_index + span - 1)
            cell = table.cell(row_index, col_index)
            if end_index > col_index:
                cell = cell.merge(table.cell(row_index, end_index))
            _write_table_cell(
                cell,
                source_cell.text,
                _table_font_size(page, row_heights[row_index]),
            )
            col_index = end_index + 1
            if col_index >= source.column_count:
                break
    return source.bbox[3]


def _set_table_geometry(
    table: Table, source: OcrTable, page: OcrPage, row_heights: Sequence[int]
) -> None:
    table.autofit = False
    left_pt = source.bbox[0] * _scale_x(page)
    width_pt = (source.bbox[2] - source.bbox[0]) * _scale_x(page)
    table._tbl.tblPr.append(parse_xml(_word_xml("tblInd", f'w:w="{_twips(left_pt)}" w:type="dxa"')))
    table._tbl.tblPr.append(parse_xml(_word_xml("tblW", f'w:w="{_twips(width_pt)}" w:type="dxa"')))

    widths = source.column_widths
    if not widths or len(widths) != source.column_count:
        equal = max(1, source.bbox[2] - source.bbox[0]) / max(1, source.column_count)
        widths = tuple(int(equal) for _ in range(source.column_count))

    for column_index, width_px in enumerate(widths):
        width = Pt(width_px * _scale_x(page))
        table.columns[column_index].width = width
        for cell in table.columns[column_index].cells:
            cell.width = width

    for row_index, row in enumerate(table.rows):
        row.height = Pt(max(10.0, row_heights[row_index] * _scale_y(page)))
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _table_row_heights(source: OcrTable) -> tuple[int, ...]:
    if source.row_heights and len(source.row_heights) == len(source.rows):
        return source.row_heights
    equal = max(1, source.bbox[3] - source.bbox[1]) / max(1, len(source.rows))
    return tuple(int(equal) for _ in source.rows)


def _table_font_size(page: OcrPage, row_height_px: int) -> float:
    height_pt = row_height_px * _scale_y(page)
    return max(8.0, min(12.0, height_pt * 0.48))


def _write_table_cell(cell: _Cell, text: str, font_size: float) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact_paragraph(paragraph, font_size)
    parts = text.splitlines() or [""]
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        _set_run_font(run, font_size)
        if index < len(parts) - 1:
            run.add_break()


def _set_cell_margins(cell: _Cell) -> None:
    margin = _TABLE_CELL_MARGIN
    cell._tc.get_or_add_tcPr().append(parse_xml(f"""
    <w:tcMar {nsdecls("w")}>
      <w:top w:w="0" w:type="dxa"/>
      <w:left w:w="{margin}" w:type="dxa"/>
      <w:bottom w:w="0" w:type="dxa"/>
      <w:right w:w="{margin}" w:type="dxa"/>
    </w:tcMar>
    """))


def _word_xml(tag: str, attrs: str) -> str:
    return f'<w:{tag} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" {attrs}/>'


def _twips(points: float) -> int:
    return int(round(points * 20))


def _add_ocr_line(document: DocxDocument, line: OcrLine, page: OcrPage, cursor_y: int) -> int:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(_space_before_pt(page, cursor_y, line.box))
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if line.box is not None and page.image_width_px:
        centered = _is_centered_line(line, page)
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.paragraph_format.left_indent = Pt(line.box[0] * _scale_x(page))
    else:
        centered = False
    run = paragraph.add_run(line.text)
    if line.box is not None and page.image_height_px:
        height_pt = _line_font_size(page, line, centered)
        _set_run_font(run, height_pt)
        paragraph.paragraph_format.line_spacing = Pt(height_pt * 1.18)
        return line.box[3]
    _set_run_font(run, 11.0)
    paragraph.paragraph_format.line_spacing = Pt(12)
    return cursor_y


def _add_vertical_gap(document: DocxDocument, page: OcrPage, cursor_y: int, top: int) -> None:
    space_before = max(0.0, (top - cursor_y) * _scale_y(page))
    if space_before <= 1.0:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.add_run("")


def _space_before_pt(page: OcrPage, cursor_y: int, box: BBox | None) -> float:
    if box is None:
        return 0.0
    gap_px = max(0, box[1] - cursor_y)
    gap_pt = gap_px * _scale_y(page)
    if cursor_y == 0:
        return gap_pt
    return max(0.0, gap_pt)


def _compact_paragraph(paragraph: Paragraph, font_size: float) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(font_size * 1.15)


def _set_run_font(run: Run, size_pt: float) -> None:
    run.font.name = _BODY_FONT
    run.font.size = Pt(size_pt)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _BODY_FONT)


def _add_seal(
    document: DocxDocument, png: bytes, bbox: BBox, page: OcrPage, index: int
) -> None:
    if not png or not page.image_width_px:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _append_vml_seal(paragraph.add_run(), page, SealImage(page.page_index, png, bbox), index)


def _configure_ocr_page(section: Section, page: OcrPage) -> None:
    if page.page_width_pt <= 0 or page.page_height_pt <= 0:
        return
    section.page_width = Pt(page.page_width_pt)
    section.page_height = Pt(page.page_height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _contains_line(box: BBox, line: OcrLine) -> bool:
    return line.box is not None and _inside(_center(line.box), box)


def _line_sort_key(line: OcrLine) -> tuple[int, int]:
    if line.box is None:
        return (0, 0)
    return (line.box[1], line.box[0])


def _center(box: BBox) -> tuple[int, int]:
    x1, y1, x2, y2 = box
    return ((x1 + x2) // 2, (y1 + y2) // 2)


def _inside(point: tuple[int, int], box: BBox) -> bool:
    x, y = point
    x1, y1, x2, y2 = box
    return x1 <= x <= x2 and y1 <= y <= y2


def _is_centered_line(line: OcrLine, page: OcrPage) -> bool:
    if line.box is None or not page.image_width_px:
        return False
    x1, y1, x2, _ = line.box
    if page.image_height_px and y1 > page.image_height_px * 0.16:
        return False
    center_delta = abs(((x1 + x2) / 2) - (page.image_width_px / 2))
    width = x2 - x1
    return center_delta <= page.image_width_px * 0.035 and (
        width >= page.image_width_px * 0.45 or width <= page.image_width_px * 0.22
    )


def _line_font_size(page: OcrPage, line: OcrLine, centered: bool) -> float:
    if line.box is None:
        return 11.0
    if centered:
        return 14.0 if line.text.startswith("《") else 15.0
    return max(11.0, min(12.0, _font_size_from_box(page, line.box)))


def _scale_x(page: OcrPage) -> float:
    return page.page_width_pt / page.image_width_px if page.image_width_px else 1.0


def _scale_y(page: OcrPage) -> float:
    return page.page_height_pt / page.image_height_px if page.image_height_px else 1.0


def _font_size_from_box(page: OcrPage, box: BBox) -> float:
    _, y1, _, y2 = box
    height_pt = (y2 - y1) * _scale_y(page)
    return max(9.0, min(16.0, height_pt * 0.78))
