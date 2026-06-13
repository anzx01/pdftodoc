"""Overlay engine：智能处理含嵌入图片的文本型 PDF。

三种模式（自动检测）
--------------------
  整页图片模式（已废弃，改为下面的"白化+文本框"模式）。

  白化+文本框模式（文字落在图片内，即"烧入图片"的合格证/表单）：
    1. 从 PDF 提取各图片块的原始字节；
    2. 用 OpenCV 将落在图片内的文字行区域涂白，消除重叠；
    3. 将清洁后的图片嵌入 DOCX 作为背景；
    4. 在原 PDF 文字坐标处放置透明浮动文本框，文字完全可编辑。

  纯叠加模式（图片为装饰背景，文字独立存在）：
    - 图片作为 behindDoc 背景；
    - 文字放在精确坐标的透明浮动文本框中。

由 ConversionService 在 has_embedded_images=True 时自动启用。
"""

import logging
import time
from io import BytesIO
from typing import NamedTuple

import cv2
import fitz
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.parser import parse_xml
from docx.shared import Pt

from pdftodoc.core.engines import CancelCheck, ProgressCallback
from pdftodoc.models.enums import ConversionStage, PdfType, TaskStatus
from pdftodoc.models.progress import ProgressEvent
from pdftodoc.models.result import ConversionResult, PageResult
from pdftodoc.models.task import ConversionTask

logger = logging.getLogger(__name__)

_EMU = 12700               # 1 pt = 12700 EMU
_TEXT_IN_IMAGE_THRESH = 0.5  # 超过此比例文字落在图片内 → 白化+文本框模式
_WIPE_PAD_X = 3            # 白化区域水平扩展像素（避免遮掉背景线条）
_WIPE_PAD_Y_TOP = 2        # 白化区域顶部扩展
_WIPE_PAD_Y_BOT = 1        # 白化区域底部扩展（保留字段下划线）


# ─── 数据结构 ─────────────────────────────────────────────────────────────────

class _Span(NamedTuple):
    text: str
    font: str
    size: float   # pt


class _Line(NamedTuple):
    x1: float
    y1: float
    x2: float
    y2: float
    spans: list[_Span]


class _ImgBlock(NamedTuple):
    x1: float     # PDF 页面坐标 (pt)
    y1: float
    x2: float
    y2: float
    raw: bytes    # 原始图片字节（PNG / JPEG）
    w_px: int     # 图片原生像素宽度
    h_px: int     # 图片原生像素高度


# ─── PDF 内容提取 ──────────────────────────────────────────────────────────────

def _extract_image_blocks(doc: "fitz.Document", page: "fitz.Page") -> list[_ImgBlock]:
    blocks: list[_ImgBlock] = []
    for b in page.get_text("dict")["blocks"]:
        if b["type"] != 1:
            continue
        x1, y1, x2, y2 = b["bbox"]
        xref = b.get("xref", 0)
        if xref > 0:
            img_data = doc.extract_image(xref)
            raw = img_data["image"]
            w_px, h_px = img_data["width"], img_data["height"]
            ext = img_data.get("ext", "png")
            if ext not in ("png", "jpeg", "jpg"):
                pix = fitz.Pixmap(raw)
                raw = pix.tobytes("png")
                w_px, h_px = pix.width, pix.height
        else:
            clip = fitz.Rect(x1, y1, x2, y2)
            pix = page.get_pixmap(clip=clip, matrix=fitz.Matrix(2, 2))
            raw = pix.tobytes("png")
            w_px, h_px = pix.width, pix.height
        blocks.append(_ImgBlock(x1, y1, x2, y2, raw, w_px, h_px))
    return blocks


def _extract_lines(page: "fitz.Page") -> list[_Line]:
    lines: list[_Line] = []
    for block in page.get_text("dict", sort=True)["blocks"]:
        if block["type"] != 0:
            continue
        for raw_line in block["lines"]:
            spans = [
                _Span(sp["text"], sp["font"], sp["size"])
                for sp in raw_line["spans"]
                if sp["text"].strip()
            ]
            if not spans:
                continue
            x1, y1, x2, y2 = raw_line["bbox"]
            lines.append(_Line(x1, y1, x2, y2, spans))
    return lines


# ─── 模式检测 ─────────────────────────────────────────────────────────────────

def _center_in_block(line: _Line, img: _ImgBlock) -> bool:
    cx = (line.x1 + line.x2) / 2
    cy = (line.y1 + line.y2) / 2
    return img.x1 <= cx <= img.x2 and img.y1 <= cy <= img.y2


def _should_wipe_mode(imgs: list[_ImgBlock], lines: list[_Line]) -> bool:
    """大多数文字行中心落在图片块内 → 用白化模式消除重叠。"""
    if not imgs or not lines:
        return False
    inside = sum(1 for ln in lines if any(_center_in_block(ln, img) for img in imgs))
    ratio = inside / len(lines)
    logger.debug("文字落在图片内占比 %.2f（阈值 %.2f）", ratio, _TEXT_IN_IMAGE_THRESH)
    return ratio >= _TEXT_IN_IMAGE_THRESH


# ─── 图片白化（消除烧入文字） ─────────────────────────────────────────────────

def _wipe_text_regions(img: _ImgBlock, lines: list[_Line]) -> bytes:
    """将落在图片区域内的文字行涂白，返回修改后的 PNG 字节。"""
    arr = np.frombuffer(img.raw, dtype=np.uint8)
    cv_img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if cv_img is None:
        return img.raw

    h, w = cv_img.shape[:2]
    blk_w_pt = img.x2 - img.x1
    blk_h_pt = img.y2 - img.y1
    if blk_w_pt <= 0 or blk_h_pt <= 0:
        return img.raw

    # PDF pt → 像素的缩放比
    sx = w / blk_w_pt
    sy = h / blk_h_pt

    for line in lines:
        if not _center_in_block(line, img):
            continue
        # 转换为图片内像素坐标（相对于图片左上角）
        px1 = max(0, int((line.x1 - img.x1) * sx) - _WIPE_PAD_X)
        py1 = max(0, int((line.y1 - img.y1) * sy) - _WIPE_PAD_Y_TOP)
        px2 = min(w, int((line.x2 - img.x1) * sx) + _WIPE_PAD_X)
        py2 = min(h, int((line.y2 - img.y1) * sy) + _WIPE_PAD_Y_BOT)
        cv_img[py1:py2, px1:px2] = 255  # 涂白

    _, buf = cv2.imencode(".png", cv_img)
    return buf.tobytes()


# ─── 字体映射 ─────────────────────────────────────────────────────────────────

def _resolve_font(raw: str) -> str:
    lo = raw.lower()
    if "yahei" in lo:
        return "Microsoft YaHei"
    if "kai" in lo:
        return "KaiTi"
    if "hei" in lo:
        return "SimHei"
    if "song" in lo or "sung" in lo:
        return "SimSun"
    return "Microsoft YaHei"


# ─── XML 工具 ────────────────────────────────────────────────────────────────

def _emu(pt: float) -> int:
    return int(round(pt * _EMU))


def _esc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ─── 背景图片（DrawingML，behindDoc） ────────────────────────────────────────

def _image_anchor_xml(rel_id: str, img: _ImgBlock, sid: int) -> str:
    x, y = _emu(img.x1), _emu(img.y1)
    cx = _emu(max(img.x2 - img.x1, 1.0))
    cy = _emu(max(img.y2 - img.y1, 1.0))
    return (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:drawing><wp:anchor'
        f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/><wp:wrapNone/>'
        f'<wp:docPr id="{sid}" name="Img{sid}"/><wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="0" name="img.png"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{rel_id}"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        f'<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>'
    )


# ─── 文本框（VML，每图片块一个，内含所有文字行段落） ───────────────────────

def _runs_xml(line: _Line) -> str:
    """将一行所有 span 拼接成 w:r 序列。"""
    parts = []
    for sp in line.spans:
        font = _resolve_font(sp.font)
        sz = max(8, int(round(sp.size * 2)))
        parts.append(
            f'<w:r>'
            f'<w:rPr>'
            f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}" w:cs="{font}"/>'
            f'<w:color w:val="auto"/>'
            f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
            f'</w:rPr>'
            f'<w:t xml:space="preserve">{_esc(sp.text)}</w:t>'
            f'</w:r>'
        )
    return "".join(parts)


def _vml_textbox_xml(img: _ImgBlock, sorted_lines: list[_Line], sid: int) -> str:
    """生成一个 VML 文本框，覆盖整个图片块区域。

    每条文字行对应一个 w:p，通过 w:before（行间距）和 w:ind w:left（左缩进）
    精确还原 PDF 中文字行的垂直与水平位置。
    """
    img_w = img.x2 - img.x1
    img_h = img.y2 - img.y1
    z = 251660000 + sid  # 保证在背景图片层之上

    paragraphs: list[str] = []
    prev_bottom = img.y1  # 上一行底部 y（初始为图片顶部）

    for line in sorted_lines:
        line_h = max(line.y2 - line.y1, 6.0)
        line_h_tw = max(120, int(round(line_h * 20)))   # 行高（twips）

        # 本行顶部到上一行底部的间距 → w:before
        gap = max(0.0, line.y1 - prev_bottom)
        before_tw = int(round(gap * 20))

        # 文字行左边缘相对图片左边缘的水平偏移 → w:ind w:left
        indent_tw = max(0, int(round((line.x1 - img.x1) * 20)))

        paragraphs.append(
            f'<w:p>'
            f'<w:pPr>'
            f'<w:spacing w:before="{before_tw}" w:after="0"'
            f' w:line="{line_h_tw}" w:lineRule="exact"/>'
            f'<w:ind w:left="{indent_tw}" w:firstLine="0"/>'
            f'<w:jc w:val="left"/>'
            f'</w:pPr>'
            f'{_runs_xml(line)}'
            f'</w:p>'
        )
        prev_bottom = line.y2

    content = "".join(paragraphs)
    return (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:pict>'
        f'<v:shape id="_x0000_s{sid}" type="#_x0000_t202"'
        f' xmlns:v="urn:schemas-microsoft-com:vml"'
        f' xmlns:o="urn:schemas-microsoft-com:office:office"'
        f' style="position:absolute;left:0pt;'
        f'margin-left:{img.x1:.4f}pt;margin-top:{img.y1:.4f}pt;'
        f'height:{img_h:.4f}pt;width:{img_w:.4f}pt;'
        f'z-index:{z};'
        f'mso-width-relative:page;mso-height-relative:page;"'
        f' filled="f" stroked="f" coordsize="21600,21600">'
        f'<v:fill on="f" focussize="0,0"/>'
        f'<v:stroke on="f"/>'
        f'<v:textbox inset="0mm,0mm,0mm,0mm">'
        f'<w:txbxContent'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'{content}'
        f'</w:txbxContent>'
        f'</v:textbox>'
        f'</v:shape>'
        f'</w:pict>'
        f'</w:r>'
    )


def _standalone_vml_textbox_xml(line: _Line, sid: int) -> str:
    """独立 VML 文本框，用于不属于任何图片块的文字行。"""
    line_h = max(line.y2 - line.y1, 6.0)
    line_h_tw = max(120, int(round(line_h * 20)))
    box_w = max(line.x2 - line.x1, 10.0) + 6.0
    box_h = line_h + 4.0
    z = 251660000 + sid
    return (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:pict>'
        f'<v:shape id="_x0000_s{sid}" type="#_x0000_t202"'
        f' xmlns:v="urn:schemas-microsoft-com:vml"'
        f' xmlns:o="urn:schemas-microsoft-com:office:office"'
        f' style="position:absolute;left:0pt;'
        f'margin-left:{line.x1:.4f}pt;margin-top:{line.y1:.4f}pt;'
        f'height:{box_h:.4f}pt;width:{box_w:.4f}pt;'
        f'z-index:{z};'
        f'mso-width-relative:page;mso-height-relative:page;"'
        f' filled="f" stroked="f" coordsize="21600,21600">'
        f'<v:fill on="f" focussize="0,0"/>'
        f'<v:stroke on="f"/>'
        f'<v:textbox inset="0mm,0mm,0mm,0mm">'
        f'<w:txbxContent'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:p>'
        f'<w:pPr>'
        f'<w:spacing w:before="0" w:after="0"'
        f' w:line="{line_h_tw}" w:lineRule="exact"/>'
        f'</w:pPr>'
        f'{_runs_xml(line)}'
        f'</w:p>'
        f'</w:txbxContent>'
        f'</v:textbox>'
        f'</v:shape>'
        f'</w:pict>'
        f'</w:r>'
    )


# ─── 页面写入 ──────────────────────────────────────────────────────────────────

def _write_page(
    document: Document,
    page: "fitz.Page",
    imgs: list[_ImgBlock],
    lines: list[_Line],
    counter: list[int],
    wipe: bool,
) -> None:
    """写入一页：白化后的背景图片 + VML 大文本框（按图片块分组）。"""
    sec = document.sections[-1]
    sec.page_width = Pt(page.rect.width)
    sec.page_height = Pt(page.rect.height)
    sec.top_margin = Pt(0)
    sec.bottom_margin = Pt(0)
    sec.left_margin = Pt(0)
    sec.right_margin = Pt(0)

    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()

    # ── 背景图片（白化后） ──────────────────────────────────────────���────────
    for img in imgs:
        counter[0] += 1
        img_bytes = _wipe_text_regions(img, lines) if wipe else img.raw
        rel_id, _ = run.part.get_or_add_image(BytesIO(img_bytes))
        run._r.append(parse_xml(_image_anchor_xml(rel_id, img, counter[0])))

    # ── 文字行按图片块分组 ────────────────────────────────────────────────────
    groups: list[list[_Line]] = [[] for _ in imgs]
    orphans: list[_Line] = []
    for line in lines:
        assigned = False
        for i, img in enumerate(imgs):
            if _center_in_block(line, img):
                groups[i].append(line)
                assigned = True
                break
        if not assigned:
            orphans.append(line)

    # ── 每个图片块生成一个 VML 大文本框 ──────────────────────────────────────
    for i, img in enumerate(imgs):
        grp = sorted(groups[i], key=lambda l: l.y1)
        if not grp:
            continue
        counter[0] += 1
        run._r.append(parse_xml(_vml_textbox_xml(img, grp, counter[0])))

    # ── 不属于任何图片块的孤立行 ─────────────────────────────────────────────
    for line in orphans:
        counter[0] += 1
        run._r.append(parse_xml(_standalone_vml_textbox_xml(line, counter[0])))


# ─── 引擎主类 ─────────────────────────────────────────────────────────────────

class OverlayEngine:
    """含嵌入图片的文本型 PDF 智能转换引擎。

    自动检测：
      - 文字已烧入图片（合格证/表单）→ 白化图片文字 + 可编辑浮动文本框
      - 图片为纯装饰背景 → 直接叠加可编辑浮动文本框
    """

    def convert(
        self,
        task: ConversionTask,
        on_progress: ProgressCallback,
        is_cancelled: CancelCheck,
    ) -> ConversionResult:
        started = time.monotonic()
        task.dst_docx.parent.mkdir(parents=True, exist_ok=True)

        doc = fitz.open(str(task.src_pdf))
        try:
            page_count = doc.page_count
            if is_cancelled():
                return _cancelled(task, page_count)

            opts = task.options
            indices = _page_indices(opts.start_page, opts.end_page, page_count)
            total = len(indices)

            # 用首页采样决定是否需要白化
            first_page = doc[indices[0]]
            sample_imgs = _extract_image_blocks(doc, first_page)
            sample_lines = _extract_lines(first_page)
            wipe = _should_wipe_mode(sample_imgs, sample_lines)
            mode = "白化+文本框" if wipe else "纯叠加文本框"
            logger.info(
                "OverlayEngine 模式：%s（图片=%d 文字行=%d）",
                mode, len(sample_imgs), len(sample_lines),
            )

            document = Document()
            counter = [0]
            page_results: list[PageResult] = []

            for step, page_index in enumerate(indices):
                if is_cancelled():
                    return _cancelled(task, page_count)

                on_progress(ProgressEvent(
                    task.task_id, ConversionStage.CONVERTING_TEXT,
                    step, total, f"第 {page_index + 1} 页",
                ))

                page = doc[page_index]
                imgs = _extract_image_blocks(doc, page)
                lines = _extract_lines(page)

                if step > 0:
                    document.add_section(WD_SECTION.NEW_PAGE)

                _write_page(document, page, imgs, lines, counter, wipe)
                page_results.append(PageResult(
                    page_index=page_index,
                    char_count=sum(len(s.text) for ln in lines for s in ln.spans),
                    line_count=len(lines),
                ))

            on_progress(ProgressEvent(
                task.task_id, ConversionStage.BUILDING_DOCX, total, total, "生成 DOCX",
            ))
            document.save(str(task.dst_docx))
            on_progress(ProgressEvent(
                task.task_id, ConversionStage.DONE, total, total, "转换完成",
            ))

            elapsed = time.monotonic() - started
            logger.info(
                "OverlayEngine 完成 %s -> %s (%d 页, %.1fs, 模式=%s)",
                task.src_pdf, task.dst_docx, total, elapsed, mode,
            )
            return ConversionResult(
                task_id=task.task_id,
                status=TaskStatus.SUCCESS,
                pdf_type=PdfType.TEXT,
                output_path=task.dst_docx,
                page_count=page_count,
                pages=tuple(page_results),
                elapsed_sec=elapsed,
                message=f"转换完成（{mode}）",
            )
        finally:
            doc.close()


# ─── 工具函数 ─────────────────────────────────────────────────────────────────

def _page_indices(start: int, end: int | None, page_count: int) -> list[int]:
    first = max(0, start)
    last = page_count - 1 if end is None else min(end, page_count - 1)
    return list(range(first, last + 1)) if first <= last else []


def _cancelled(task: ConversionTask, page_count: int) -> ConversionResult:
    return ConversionResult(
        task_id=task.task_id,
        status=TaskStatus.CANCELLED,
        pdf_type=PdfType.TEXT,
        output_path=None,
        page_count=page_count,
        message="已取消",
    )


__all__ = ["OverlayEngine"]
