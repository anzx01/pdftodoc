"""OCR 链路测试：注入 fake 识别器，对扫描型 PDF 跑端到端，不依赖真实 paddle。

覆盖 renderer（真实 PyMuPDF 渲染）、docx_builder（真实 python-docx）、
ocr_engine（分页/取消/结果聚合）以及 service → OCR 引擎的分派整链路。
"""

from pathlib import Path
from zipfile import ZipFile

import fitz
import numpy as np
from docx import Document

from pdftodoc.core.engines.ocr_engine import OcrEngine
from pdftodoc.core.ocr import OcrCell, OcrLine, OcrPage, OcrTable, PageImage, SealImage
from pdftodoc.core.ocr.docx_builder import build_docx, build_image_docx
from pdftodoc.core.ocr.renderer import render_page, render_page_image
from pdftodoc.core.service import ConversionService
from pdftodoc.models.enums import PdfType, TaskStatus
from pdftodoc.models.task import ConversionOptions, ConversionTask


class FakeRecognizer:
    """计数用假识别器：每页返回一行固定文本，不触碰 paddle。"""

    def __init__(self) -> None:
        self.calls = 0

    def recognize(self, image: np.ndarray) -> tuple[str, ...]:
        self.calls += 1
        return (f"第{self.calls}页识别文本",)


class CapturingRecognizer:
    """记录引擎传入的图片均值，验证 OCR 前处理是否生效。"""

    def __init__(self) -> None:
        self.means: list[float] = []

    def recognize(self, image: np.ndarray) -> tuple[str, ...]:
        self.means.append(float(image.mean()))
        return ("识别文本",)


def test_render_page_shape(scanned_pdf: Path) -> None:
    doc = fitz.open(scanned_pdf)
    try:
        img = render_page(doc, 0, dpi=100)
    finally:
        doc.close()
    assert img.ndim == 3 and img.shape[2] == 3
    assert img.dtype == np.uint8


def test_render_page_image(scanned_pdf: Path) -> None:
    doc = fitz.open(scanned_pdf)
    try:
        page_image = render_page_image(doc, 0, dpi=100)
    finally:
        doc.close()
    assert page_image.png.startswith(b"\x89PNG")
    assert page_image.width_pt > 0
    assert page_image.height_pt > 0


def test_build_docx_writes_lines(tmp_path: Path) -> None:
    dst = tmp_path / "b.docx"
    build_docx((OcrPage(0, ("行一", "行二")), OcrPage(1, ())), dst)
    assert dst.exists()
    texts = [p.text for p in Document(str(dst)).paragraphs]
    assert "行一" in texts and "行二" in texts


def test_build_image_docx_writes_full_page_image(tmp_path: Path) -> None:
    dst = tmp_path / "image.docx"
    png = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
        b"\x00\x00\x03\x01\x01\x00\xc9\xfe\x92\xef\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    build_image_docx((PageImage(0, png, 72.0, 72.0),), dst)

    doc = Document(str(dst))
    assert len(doc.inline_shapes) == 1
    assert doc.sections[0].page_width.pt == 72.0
    assert doc.sections[0].page_height.pt == 72.0


def test_build_docx_writes_detected_table(tmp_path: Path) -> None:
    dst = tmp_path / "table.docx"
    page = OcrPage(
        page_index=0,
        lines=("标题", "表头", "值一", "值二"),
        text_lines=(
            OcrLine("标题", (10, 5, 60, 20)),
            OcrLine("表头", (10, 35, 60, 50)),
            OcrLine("值一", (10, 75, 50, 90)),
            OcrLine("值二", (80, 75, 120, 90)),
        ),
        tables=(
            OcrTable(
                bbox=(0, 30, 140, 100),
                column_count=2,
                rows=((OcrCell("表头", col_span=2),), (OcrCell("值一"), OcrCell("值二"))),
            ),
        ),
    )

    build_docx((page,), dst)

    doc = Document(str(dst))
    assert [p.text for p in doc.paragraphs if p.text.strip()] == ["标题"]
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "表头"
    assert doc.tables[0].rows[1].cells[0].text == "值一"
    assert doc.tables[0].rows[1].cells[1].text == "值二"


def test_build_docx_writes_positioned_text_layer(tmp_path: Path) -> None:
    dst = tmp_path / "positioned.docx"
    png = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
        b"\x00\x00\x03\x01\x01\x00\xc9\xfe\x92\xef\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    page = OcrPage(
        page_index=0,
        lines=("可编辑文字",),
        text_lines=(OcrLine("可编辑文字", (10, 20, 90, 40)),),
        image_width_px=100,
        image_height_px=100,
        page_width_pt=100.0,
        page_height_pt=100.0,
        seals=(SealImage(0, png, (40, 50, 70, 80)),),
    )

    build_docx((page,), dst)

    with ZipFile(dst) as docx:
        xml = docx.read("word/document.xml").decode("utf-8")
    assert "page_bg_0" not in xml
    assert "seal_0_0" in xml
    assert "ocr_text_0_0" not in xml
    assert "v:textbox" not in xml
    assert "txbxContent" not in xml
    assert "可编辑文字" in xml
    assert "mso-position-horizontal-relative:page" in xml
    assert "w:vanish" not in xml
    doc = Document(str(dst))
    assert any(paragraph.text.strip() for paragraph in doc.paragraphs)


def test_ocr_engine_end_to_end(scanned_pdf: Path, tmp_path: Path) -> None:
    dst = tmp_path / "out.docx"
    fake = FakeRecognizer()
    result = OcrEngine(recognizer=fake).convert(
        ConversionTask(
            scanned_pdf, dst, ConversionOptions(preserve_scan_layout=False)
        ),
        lambda _: None,
        lambda: False,
    )
    assert result.status is TaskStatus.SUCCESS
    assert result.pdf_type is PdfType.SCANNED
    assert dst.exists() and dst.stat().st_size > 0
    assert result.page_count == 2
    assert fake.calls == 2
    assert len(result.pages) == 2
    assert result.pages[0].line_count == 1


def test_ocr_engine_wipes_light_watermark_before_recognition(
    scanned_pdf: Path, tmp_path: Path
) -> None:
    dst = tmp_path / "out.docx"
    recognizer = CapturingRecognizer()

    result = OcrEngine(recognizer=recognizer).convert(
        ConversionTask(scanned_pdf, dst, ConversionOptions(preserve_scan_layout=False)),
        lambda _: None,
        lambda: False,
    )

    assert result.status is TaskStatus.SUCCESS
    assert recognizer.means
    assert min(recognizer.means) > 250.0


def test_ocr_engine_cancel(scanned_pdf: Path, tmp_path: Path) -> None:
    dst = tmp_path / "out.docx"
    fake = FakeRecognizer()
    result = OcrEngine(recognizer=fake).convert(
        ConversionTask(scanned_pdf, dst), lambda _: None, lambda: True
    )
    assert result.status is TaskStatus.CANCELLED
    assert result.output_path is None
    assert fake.calls == 0
    assert not dst.exists()


def test_ocr_engine_preserves_scan_layout_when_requested(scanned_pdf: Path, tmp_path: Path) -> None:
    dst = tmp_path / "layout.docx"
    fake = FakeRecognizer()

    result = OcrEngine(recognizer=fake).convert(
        ConversionTask(scanned_pdf, dst, ConversionOptions(preserve_scan_layout=True)),
        lambda _: None,
        lambda: False,
    )

    assert result.status is TaskStatus.SUCCESS
    assert result.pdf_type is PdfType.SCANNED
    assert fake.calls == 0
    assert dst.exists() and dst.stat().st_size > 0
    assert len(Document(str(dst)).inline_shapes) == 2


def test_service_dispatches_scanned_end_to_end(scanned_pdf: Path, tmp_path: Path) -> None:
    """detector 真把扫描件判为 SCANNED，且 service → OcrEngine → DOCX 全链路打通。"""
    dst = tmp_path / "s.docx"
    svc = ConversionService()
    svc._ocr_engine = OcrEngine(recognizer=FakeRecognizer())  # 跳过懒加载，避免拉起 paddle
    result = svc.convert(
        ConversionTask(scanned_pdf, dst, ConversionOptions(preserve_scan_layout=False))
    )
    assert result.status is TaskStatus.SUCCESS
    assert result.pdf_type is PdfType.SCANNED
    assert dst.exists() and dst.stat().st_size > 0
